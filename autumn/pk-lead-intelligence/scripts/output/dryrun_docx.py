"""Dry-run docx writer.

Emits a `RenderedNote` as a local `.docx` file for operator review
before any Salesforce writes are allowed. The dry-run path is the
only artifact the operator inspects during the 5-Note review loop in
SKILL.md "Pre-Run Checklist", so the file must (a) parse as a real
.docx and (b) carry every locked-layout section.

Pure I/O wrapper — does not synthesize content, does not branch on
section headings. Roundtrip behavior is pinned by
`tests/test_dryrun_docx.py`.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from scripts.output.note_renderer import RenderedNote


def write(*, note: RenderedNote, output_path: Path) -> Path:
    """Write `note` as a .docx at `output_path`. Returns the path written.

    Creates the parent directory if it does not exist; the gitignore
    blocks `/output/` and `*.docx`, so a dry-run run lands files that
    git correctly ignores.
    """

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading(note.title, level=0)
    document.add_paragraph(f"Enriched at {note.enriched_at_utc} (UTC)")
    for section in note.sections:
        document.add_heading(section.heading, level=1)
        document.add_paragraph(section.body)
    document.save(str(output_path))
    return output_path
